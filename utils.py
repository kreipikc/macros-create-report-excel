import os
from docx import Document
from typing import Optional, List
from num2words import num2words
from pydantic import ValidationError
from schemas import ChecksDefault, TypeCheck, TypeDocument


def validate_check(check: ChecksDefault) -> None:
    """Validate check to ensure all required fields are filled based on the check type.

    Args:
        check (ChecksDefault): A list of ChecksDefault objects to be validated.

    Returns:
        None

    Raises:
        Exception: If a required field is missing for a specific check type.
    """
    if not check.sum_check or not check.id_check or not check.date:
        raise Exception("Не заполнены реквизиты чека!")

    # Проверка определенного типа
    if check.type == TypeCheck.representative_offices_event:
        if not check.counterparty:
            raise Exception(f"Не заполнено поле 'Контрагент' (Чек ID: {check.id_check})")

        if not check.counterparty_participant:
            raise Exception(f"Не заполнено поле 'Участник контрагента' (Чек ID: {check.id_check})")

        if not check.counterparty_post:
            raise Exception(f"Не заполнено поле 'Контрагент должность' (Чек ID: {check.id_check})")

        if not check.meeting_place:
            raise Exception(f"Не заполнено поле 'Место встречи' (Чек ID: {check.id_check})")
    elif check.type == TypeCheck.representative_offices_present:
        if not check.topic:
            raise Exception(f"Не заполнено поле 'Тема / Мероприятие' (Чек ID: {check.id_check})")

        if not check.counterparty:
            raise Exception(f"Не заполнено поле 'Контрагент' (Чек ID: {check.id_check})")

        if not check.counterparty_participant:
            raise Exception(f"Не заполнено поле 'Участник контрагента' (Чек ID: {check.id_check})")

        if not check.name_present:
            raise Exception(f"Не заполнено поле 'Наименование подарка' (Чек ID: {check.id_check})")
    elif check.type == TypeCheck.round_table_discussion_Club:
        if not check.medication:
            raise Exception(f"Не заполнено поле 'Препарат' (Чек ID: {check.id_check})")

        if not check.counterparty_participant:
            raise Exception(f"Не заполнено поле 'Участник контрагента' (Чек ID: {check.id_check})")

        if not check.counterparty_post:
            raise Exception(f"Не заполнено поле 'Контрагент должность' (Чек ID: {check.id_check})")

        if not check.topic:
            raise Exception(f"Не заполнено поле 'Тема / Мероприятие' (Чек ID: {check.id_check})")

        if not check.meeting_place:
            raise Exception(f"Не заполнено поле 'Место встречи' (Чек ID: {check.id_check})")


def create_check(check: tuple) -> Optional[ChecksDefault]:
    """Create a ChecksDefault object from a tuple of check data.

    Args:
        check (tuple): A tuple containing check data

    Returns:
        Optional[ChecksDefault]: A ChecksDefault object if the check data is valid, otherwise None.

    Raises:
        ValidationError: If the check data is invalid.
        Exception: If any other error occurs during the creation of the check.
    """
    try:
        check_res = ChecksDefault(
            number_str=int(check[0]),
            type_document=TypeDocument(str(check[1].strip())),
            id_check=check[2],
            date=check[3],
            sum_check=float(check[4]),
            type=TypeCheck(check[5].lower().strip()),
            counterparty=check[6],
            counterparty_participant=check[7],
            counterparty_post=check[8],
            meeting_place=check[9],
            medication=check[10],
            topic=check[11],
            name_present=check[12],
            comment=check[13],
        )
        return check_res
    except ValidationError as e:
        print(f"Validation error: {e}")
        return None
    except Exception as e:
        print(f"Error: {e}")
        return None


def replace_text_in_paragraph(paragraph, key, value) -> None:
    """Replace a specific key with a value in a paragraph.

    Args:
        paragraph (Paragraph): The paragraph object where the replacement will occur.
        key (str): The placeholder text to be replaced.
        value (str): The text to replace the placeholder.

    Returns:
        None
    """
    if key in paragraph.text:
        full_text = ''.join(run.text for run in paragraph.runs)

        if key in full_text:
            full_text = full_text.replace(key, value)

            for run in paragraph.runs:
                run.text = ""

            paragraph.runs[0].text = full_text


def replace_text_in_table(table, key, value) -> None:
    """Replace a specific key with a value in all cells of a table.

    Args:
        table (Table): The table object where the replacement will occur.
        key (str): The placeholder text to be replaced.
        value (str): The text to replace the placeholder.

    Returns:
        None
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, key, value)


def create_representative_word(replacements: dict, file_template: str, output_path: str) -> None:
    """Generate a Word document by replacing placeholders in a template.

    Args:
        replacements (dict): A dictionary of placeholders and their corresponding values.
        file_template (str): The name of the template file located in the 'templates' directory.
        output_path (str): The path where the generated document will be saved.

    Returns:
        None
    """
    doc = Document(f"templates/{file_template}")

    for key, value in replacements.items():
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, key, value)

        for table in doc.tables:
            replace_text_in_table(table, key, value)

    doc.save(f"{output_path}")


# def convert_excel_date_to_normal(excel_date: int) -> date:
#     """Convert an Excel date to a normal date.
#
#     Args:
#         excel_date (int): The date in Excel format.
#
#     Returns:
#         date: The converted date in normal format.
#     """
#     return START_DATE + timedelta(days=excel_date)


def sum_money_all_checks(checks: List[ChecksDefault]) -> float:
    """Calculate the total sum of all checks.

    Args:
        checks (List[ChecksDefault]): A list of ChecksDefault objects.

    Returns:
        float: The total sum of all checks.
    """
    sum_checks = 0.0
    for check in checks:
        sum_checks += check.sum_check
    return sum_checks


def convert_num_to_word(num: int) -> str:
    """Convert a number to its word representation in Russian.

    Args:
        num (int): The number to convert.

    Returns:
        A str, the number in words, capitalized.
    """
    num_word = num2words(num, lang='ru')
    return num_word.capitalize()


def create_kopecks_str(num: float) -> str:
    """Extract the kopecks part from a monetary amount.

    Args:
        num (float): The monetary amount in rubles and kopecks.

    Returns:
        A str, the kopecks part as a string (e.g., "00" or "50").
    """
    kopecks = (num - int(num)) * 100
    if kopecks == 0:
        return "00"
    else:
        return f"{kopecks:.0f}"


def create_text_price(rubles: int, kopecks: int) -> str:
    """Convert a monetary amount to words in Russian.

    Args:
        rubles (int): The amount in rubles.
        kopecks (int): The amount in kopecks.

    Returns:
        str: The monetary amount in words.
    """
    rubles_in_words = convert_num_to_word(rubles)
    # kopecks_in_words = convert_num_to_word(kopecks)

    result = f"{rubles_in_words} рублей {kopecks} копеек ({rubles} руб. {kopecks} коп.)"
    return result


def get_absolute_path(relative_path: str) -> str:
    """Get the absolute path based on the location of the script.

    Args:
        relative_path (str): The relative path to the file.

    Returns:
        str: The absolute path to the file.
    """
    script_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(script_dir, relative_path)